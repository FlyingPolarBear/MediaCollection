'''
Author: Derry
Date: 2022-08-07 22:45:57
LastEditors: Derry
Email: drlv@mail.ustc.edu.cn
LastEditTime: 2022-12-06 20:01:27
Description: None
'''
import datetime
import os
import time

import docx
from bs4 import BeautifulSoup
from HTMLTable import HTMLTable

from classify_model import TextClassifier
from email_sender import EmailSender
from google_translate import GoogleTranslate


class Cell:
    def __init__(self):
        self.base_url = "https://www.sciencedirect.com/journal/molecular-cell/articles-in-press"
        # self.email_address = "derrylv@qq.com"
        self.email_address = "411473510@qq.com"
        self.email_sender = EmailSender(msg_to=self.email_address)
        self.translator = GoogleTranslate()

    def _nextpage(self, i):
        return f"https://www.sciencedirect.com/journal/molecular-cell/articles-in-press?page={i}"

    def load_main_webpage(self, url, file_path="tmp/articles-in-press"):
        def load_webpage():
            if os.path.exists(file_path):
                os.remove(file_path)
            os.system(
                f'wget -U "Mozilla/5.0" -O {file_path} "{url}"')
            with open(file_path, "rb") as f:
                html = f.read()
            return html

        while True:
            try:
                html = load_webpage()
                soup = BeautifulSoup(html, 'lxml')
                break
            except:
                pass
        return soup

    def extract_date_info(self, raw_date):
        raw_date = raw_date.lstrip(
            "In Press, Corrected Proof, Available online ")
        date = datetime.datetime.strptime(raw_date, '%d %B %Y')
        return date

    def get_detail_info(self, url, paper_name):
        soup = self.load_main_webpage(
            url, file_path=f"tmp/{paper_name}.html")
        summary = soup.find('div', class_='abstract author').get_text()[7:]
        keywords_list = soup.find_all('div', class_='keyword')
        keywords = [keyword.text for keyword in keywords_list]
        keywords = ", ".join(keywords)
        return summary, keywords

    def get_info(self, time_interval=None):
        tc = TextClassifier()
        if time_interval is None:  # 默认取最近七天的文章
            self.end_date = datetime.datetime.today()
            self.start_date = datetime.datetime.now() - datetime.timedelta(days=7)
        else:  # 指定起止日期
            self.start_date, self.end_date = time_interval
            self.start_date = datetime.datetime.strptime(
                self.start_date, '%Y-%m-%d')
            self.end_date = datetime.datetime.strptime(
                self.end_date, '%Y-%m-%d')

        self.date_interval = f"{self.start_date.strftime('%Y-%m-%d')} - {self.end_date.strftime('%Y-%m-%d')}"
        self.caption = f'Molecular Cell Weekly Update ({self.date_interval})'

        i = 0
        while True:
            i += 1
            soup = self.load_main_webpage(self._nextpage(
                i), file_path="tmp/articles-in-press")
            info_list = soup.find_all(
                'dl', class_='js-article article-content')
            paper_list = []
            for paper_soup in info_list:
                paper = {}

                raw_date = paper_soup.find(
                    'dd', class_='u-clr-grey8 u-text-italic text-s js-article-item-aip-date').get_text()
                paper['date'] = self.extract_date_info(raw_date)
                if paper['date'] <= self.start_date or paper['date'] >= self.end_date:
                    break

                title_element = paper_soup.find(
                    'a', class_='anchor article-content-title u-margin-xs-top u-margin-s-bottom anchor-default')
                relative_url = title_element.get('href')
                paper['title'] = title_element.get_text()
                paper['url'] = f"https://www.sciencedirect.com{relative_url}"
                paper['author'] = paper_soup.find(
                    'dd', class_='js-article-author-list u-clr-grey8 text-s').get_text()

                paper["summary"], paper["keywords"] = self.get_detail_info(
                    paper['url'], paper['title'].split(" ")[0])
                for totrans_key in ['title', 'summary', 'keywords']:
                    paper[f"{totrans_key}_zh"] = self.translator.translate_en2zh(
                        paper[totrans_key])

                paper['is_need'] = "否"
                for key in tc.focus_words:
                    if key in paper['keywords'] or key in paper['title']:
                        paper['is_need'] = "是"
                        break
                if paper['is_need'] == "否" and tc.predict([paper['title']])[0] == 1:
                    paper['is_need'] = "是"

                paper['date'] = paper['date'].strftime('%Y-%m-%d')

                print(paper)

                paper_list.append(paper)

            self.save_to_word(paper_list, "tmp/res.docx")
            return paper_list

    def send_info(self, paper_list):
        for paper in paper_list:
            for key in ['title', 'summary', 'keywords', 'title_zh', 'summary_zh', 'keywords_zh']:
                paper[key] = auto_wrap(str(paper[key]))
        self.url_caption = f'<a href="{self.base_url}">Molecular Cell</a> Weekly Update ({self.date_interval})'

        table = HTMLTable(caption=self.url_caption)
        table.append_header_rows(
            [["相关", "标题", "日期", "摘要", "中文摘要", "关键词"]])
        for paper in paper_list:
            title = f'<a href="{paper["url"]}">' + paper['title']+'</a>'
            table.append_data_rows([[paper['is_need'], title+"++++"+paper['title_zh'], paper['date'],
                                   paper['summary'], paper['summary_zh'], paper['keywords']+"++++"+paper['keywords_zh']]])
        # 标题样式
        table.caption.set_style({
            'font-size': '30px',
            'font-weight': 'bold',
            'text-align': 'center',
            'margin': '30px',
        })
        # 表格样式，即<table>标签样式
        table.set_style({
            'border-collapse': 'collapse',
            'white-space': 'nowrap',
            'font-size': '14px',
            'margin': '0 auto',
        })
        # 统一设置所有单元格样式，<td>或<th>
        table.set_cell_style({
            'border': '1px solid #000',
            'width': '100px',
            'padding': '5px',
            'align': 'center',
            'text-align': 'left',
            'margin': '0 auto',
            'word-break': 'break-all',
            'word-wrap': 'break-word',
        })
        # 表头样式
        table.set_header_row_style({
            'color': '#fff',
            'background-color': '#007dbc',
            'font-size': '18px',
        })
        # 覆盖表头单元格字体样式
        table.set_header_cell_style({
            'padding': '15px',
        })
        for row in table.iter_data_rows():
            if row[0].value == "是":
                row.set_style({
                    'background-color': '#ffdddd'
                })
        html = table.to_html()
        html = html.replace("++++", "<br>").replace("&lt;",
                                                    "<").replace("&gt;", ">").replace("&quot;", '"')
        with open("tmp/res.html", "w") as f:
            f.write(html)

        self.email_sender.send(
            content=html, subject=self.caption, attachment="tmp/res.docx")

    def save_to_word(self, paper_list, doc_path="tmp/res.docx"):
        doc = docx.Document()
        doc.add_heading(self.caption, level=0)
        for paper in paper_list:
            if paper['is_need'] == "否":
                continue
            doc.add_heading(f"{paper['title']}({paper['title_zh']})", level=1)
            doc.add_paragraph("日期："+paper['date'])
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(paper['summary'])
            doc.add_paragraph(paper['summary_zh'])
            doc.add_heading("Keywords", level=2)
            doc.add_paragraph(paper['keywords'])
            doc.add_paragraph(paper['keywords_zh'])
            doc.add_paragraph("")
        doc.save(doc_path)


def auto_wrap(text, width_limit=60):
    """
    对于给定的中英文字符串，按照指定的宽度进行自动换行（添加换行符++++）
    """
    def is_chinese(uchar):
        # 判断是否是中文字符
        if uchar >= u'\u4e00' and uchar <= u'\u9fa5':
            return True
        else:
            return False
    new_text = ""
    for char in text:
        if is_chinese(char):
            new_text += ' '+char+' '
        else:
            new_text += char
    new_text = new_text.split()
    width = 0
    res = ""
    for word in new_text:
        word_width = len(word)
        if width + word_width <= width_limit:
            res += word
            width += word_width+1
        else:
            res.rstrip()
            res += "++++"+word
            width = word_width+1
        if not is_chinese(word):
            res += ' '
    return res.rstrip()


def CellWeekly(immed=True):
    current_time = datetime.datetime.now()
    print(f"{current_time.strftime('%Y-%m-%d %H:%M:%S')}\t脚本开始运行")
    if immed:
        cell = Cell()
        paper_list = cell.get_info()
        cell.send_info(paper_list)
    next_week_time = current_time + \
        datetime.timedelta(days=(4-current_time.weekday()) % 7)
    next_week_time = next_week_time.replace(hour=16, minute=0, second=0)
    if next_week_time < current_time:
        next_week_time += datetime.timedelta(days=7)
    sleep_time = (next_week_time - current_time).total_seconds()
    print(f"{current_time.strftime('%Y-%m-%d %H:%M:%S')}\t下次唤醒时间：{next_week_time.strftime('%Y-%m-%d %H:%M:%S')}")
    time.sleep(sleep_time)
    while True:
        current_time = datetime.datetime.now()
        cell = Cell()
        paper_list = cell.get_info()
        cell.send_info(paper_list)
        next_week_time = current_time + datetime.timedelta(days=7)
        next_week_time = next_week_time.replace(hour=20, minute=0, second=0)
        sleep_time = (next_week_time - current_time).total_seconds()
        print(f"{current_time.strftime('%Y-%m-%d %H:%M:%S')}\t下次唤醒时间：{next_week_time.strftime('%Y-%m-%d %H:%M:%S')}")
        time.sleep(sleep_time)


if __name__ == "__main__":
    # CellWeekly(immed=True)
    cell = Cell()
    # paper_list = cell.get_info(["2022-11-1", "2022-12-5"])
    paper_list = cell.get_info()
    cell.send_info(paper_list)
